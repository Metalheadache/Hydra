"""
Structured file reader tools for Hydra agents.

These complement the existing FileProcessor (which does flattened text
extraction for pipeline context injection) by providing *structured*
programmatic reading — tables as list-of-dicts, heading hierarchies,
column statistics, code structure maps — that downstream agents can
actually reason over.

Dependencies:
    python-docx  (already in core deps)
    openpyxl     (already in core deps)
    csv, pathlib (stdlib)
"""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import TYPE_CHECKING

import structlog

from hydra.models import ToolResult
from hydra.tools._security import safe_read_path
from hydra.tools.base import BaseTool

if TYPE_CHECKING:
    from hydra.state_manager import StateManager

logger = structlog.get_logger(__name__)


# ── read_docx ─────────────────────────────────────────────────────────────────


class ReadDocxTool(BaseTool):
    """Read a Word (.docx) document with structured extraction."""

    name = "read_docx"
    description = (
        "Read a Word (.docx) document. Extracts text with heading structure, "
        "tables as lists-of-dicts, and document metadata (author, created date, etc). "
        "Use for ingesting reports, contracts, specs, or any Word document."
    )
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "extract": {
                "type": "string",
                "enum": ["all", "text_only", "tables_only", "metadata_only"],
                "description": "What to extract. Default: all.",
                "default": "all",
            },
            "max_chars": {
                "type": "integer",
                "description": "Truncate text output at this many characters. Default: 50000.",
                "default": 50000,
            },
        },
        "required": ["file_path"],
    }

    def __init__(self, output_dir: str = "./hydra_output", state_manager: "StateManager | None" = None) -> None:
        self._output_dir = output_dir
        self._state_manager = state_manager

    async def execute(
        self, file_path: str, extract: str = "all", max_chars: int = 50000
    ) -> ToolResult:
        try:
            from docx import Document

            path = safe_read_path(file_path, allowed_roots=[self._output_dir, Path.cwd()])

            doc = Document(str(path))
            result: dict = {}

            if extract in ("all", "text_only"):
                paragraphs = []
                for para in doc.paragraphs:
                    prefix = ""
                    if para.style and para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading ", "").strip()
                        prefix = f"[H{level}] "
                    paragraphs.append(f"{prefix}{para.text}")
                full_text = "\n".join(paragraphs)
                result["text"] = full_text[:max_chars]
                result["truncated"] = len(full_text) > max_chars
                result["total_paragraphs"] = len(doc.paragraphs)

            if extract in ("all", "tables_only"):
                tables = []
                for i, table in enumerate(doc.tables):
                    if not table.rows:
                        tables.append({"index": i, "headers": [], "rows": [], "row_count": 0})
                        continue
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    rows = []
                    for row in table.rows[1:]:
                        row_data = {}
                        for j, cell in enumerate(row.cells):
                            key = headers[j] if j < len(headers) else f"col_{j}"
                            row_data[key] = cell.text.strip()
                        rows.append(row_data)
                    tables.append({"index": i, "headers": headers, "rows": rows, "row_count": len(rows)})
                result["tables"] = tables

            if extract in ("all", "metadata_only"):
                props = doc.core_properties
                result["metadata"] = {
                    "author": props.author,
                    "title": props.title,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                    "last_modified_by": props.last_modified_by,
                    "revision": props.revision,
                }

            logger.info("read_docx_success", file_path=file_path, extract=extract)
            return ToolResult(success=True, data=result)

        except ValueError as exc:
            return ToolResult(success=False, error=str(exc))
        except Exception as exc:
            logger.error("read_docx_failed", file_path=file_path, error=str(exc))
            return ToolResult(success=False, error=f"Failed to read DOCX: {exc}")


# ── read_xlsx ─────────────────────────────────────────────────────────────────


class ReadXlsxTool(BaseTool):
    """Read an Excel workbook with structured extraction and column statistics."""

    name = "read_xlsx"
    description = (
        "Read an Excel (.xlsx/.xls) file. Returns sheet names, row data as "
        "list-of-dicts (header row as keys), and column stats (type inference, "
        "min/max/mean, sample values). Handles multi-sheet workbooks."
    )
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the Excel file.",
            },
            "sheet": {
                "type": "string",
                "description": "Sheet name. Default: first sheet. Use '__all__' for all sheets.",
                "default": "",
            },
            "max_rows": {
                "type": "integer",
                "description": "Max rows per sheet. Default: 500. Use -1 for all.",
                "default": 500,
            },
            "header_row": {
                "type": "integer",
                "description": "Row index (1-based) containing headers. Default: 1.",
                "default": 1,
            },
            "include_stats": {
                "type": "boolean",
                "description": "Include column type inference and basic stats. Default: true.",
                "default": True,
            },
        },
        "required": ["file_path"],
    }

    def __init__(self, output_dir: str = "./hydra_output", state_manager: "StateManager | None" = None) -> None:
        self._output_dir = output_dir
        self._state_manager = state_manager

    async def execute(
        self,
        file_path: str,
        sheet: str = "",
        max_rows: int = 500,
        header_row: int = 1,
        include_stats: bool = True,
    ) -> ToolResult:
        try:
            from openpyxl import load_workbook

            path = safe_read_path(file_path, allowed_roots=[self._output_dir, Path.cwd()])

            wb = load_workbook(str(path), read_only=True, data_only=True)
            result: dict = {"sheets": wb.sheetnames}

            sheets_to_read = wb.sheetnames if sheet == "__all__" else [sheet or wb.sheetnames[0]]
            sheets_data = {}

            for sname in sheets_to_read:
                if sname not in wb.sheetnames:
                    continue
                ws = wb[sname]
                # Stream rows — only keep max_rows in memory, count total without materializing
                header_cells = None
                records = []
                total = 0
                for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    if row_idx == header_row - 1:
                        header_cells = row
                        continue
                    if row_idx < header_row:
                        continue
                    total += 1
                    if max_rows > 0 and len(records) >= max_rows:
                        continue  # keep counting total but stop collecting records

                    record = {}
                    if header_cells is None:
                        continue
                    hdr = [str(c) if c is not None else f"col_{i}" for i, c in enumerate(header_cells)]
                    for j, val in enumerate(row):
                        key = hdr[j] if j < len(hdr) else f"col_{j}"
                        if hasattr(val, "isoformat"):
                            val = val.isoformat()
                        record[key] = val
                    records.append(record)

                if header_cells is None:
                    sheets_data[sname] = {"headers": [], "rows": [], "total_rows": 0}
                    continue

                headers = [str(c) if c is not None else f"col_{i}" for i, c in enumerate(header_cells)]

                sheet_result: dict = {
                    "headers": headers,
                    "rows": records,
                    "total_rows": total,
                    "truncated": total > max_rows > 0,
                }

                if include_stats and records:
                    col_stats = {}
                    for h in headers:
                        values = [r.get(h) for r in records if r.get(h) is not None]
                        types = set(type(v).__name__ for v in values)
                        stat: dict = {"types": list(types), "non_null": len(values)}
                        nums = [v for v in values if isinstance(v, (int, float))]
                        if nums:
                            stat.update({
                                "min": min(nums),
                                "max": max(nums),
                                "mean": round(sum(nums) / len(nums), 2),
                            })
                        stat["sample_values"] = [str(v) for v in values[:3]]
                        col_stats[h] = stat
                    sheet_result["column_stats"] = col_stats

                sheets_data[sname] = sheet_result

            wb.close()
            result["data"] = sheets_data

            logger.info("read_xlsx_success", file_path=file_path, sheets_read=len(sheets_data))
            return ToolResult(success=True, data=result)

        except ValueError as exc:
            return ToolResult(success=False, error=str(exc))
        except Exception as exc:
            logger.error("read_xlsx_failed", file_path=file_path, error=str(exc))
            return ToolResult(success=False, error=f"Failed to read Excel file: {exc}")


# ── read_csv ──────────────────────────────────────────────────────────────────


class ReadCsvTool(BaseTool):
    """Read CSV/TSV with encoding detection and delimiter inference."""

    name = "read_csv"
    description = (
        "Read a CSV or TSV file. Auto-detects delimiter and encoding "
        "(utf-8, gb18030, latin-1). Returns headers, rows as list-of-dicts, "
        "and basic column info."
    )
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the CSV/TSV file.",
            },
            "delimiter": {
                "type": "string",
                "description": "Column delimiter. Default: auto-detect.",
                "default": "",
            },
            "encoding": {
                "type": "string",
                "description": "File encoding. Default: auto-detect (utf-8, gb18030, latin-1).",
                "default": "",
            },
            "max_rows": {
                "type": "integer",
                "description": "Max rows to return. Default: 1000.",
                "default": 1000,
            },
            "has_header": {
                "type": "boolean",
                "description": "First row is header. Default: true.",
                "default": True,
            },
        },
        "required": ["file_path"],
    }

    # Encoding cascade — utf-8-sig first (handles BOM), gb18030 is the
    # CJK superset that covers gbk/gb2312 edge cases.
    _ENCODING_CASCADE = ("utf-8-sig", "utf-8", "gb18030", "latin-1")

    def __init__(self, output_dir: str = "./hydra_output", state_manager: "StateManager | None" = None) -> None:
        self._output_dir = output_dir
        self._state_manager = state_manager

    async def execute(
        self,
        file_path: str,
        delimiter: str = "",
        encoding: str = "",
        max_rows: int = 1000,
        has_header: bool = True,
    ) -> ToolResult:
        try:
            path = safe_read_path(file_path, allowed_roots=[self._output_dir, Path.cwd()])
            raw = path.read_bytes()

            # Encoding detection
            detected_encoding = encoding
            if not detected_encoding:
                for enc in self._ENCODING_CASCADE:
                    try:
                        raw.decode(enc)
                        detected_encoding = enc
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue
                else:
                    return ToolResult(success=False, error="Could not detect file encoding")

            text = raw.decode(detected_encoding)

            # Delimiter detection
            detected_delimiter = delimiter
            if not detected_delimiter:
                try:
                    dialect = csv.Sniffer().sniff(text[:8192])
                    detected_delimiter = dialect.delimiter
                except csv.Error:
                    detected_delimiter = "," if path.suffix.lower() == ".csv" else "\t"

            reader = csv.reader(text.splitlines(), delimiter=detected_delimiter)
            all_rows = list(reader)

            if not all_rows:
                return ToolResult(success=True, data={"headers": [], "rows": [], "total_rows": 0})

            headers = all_rows[0] if has_header else [f"col_{i}" for i in range(len(all_rows[0]))]
            data_rows = all_rows[1:] if has_header else all_rows
            total = len(data_rows)

            records = []
            for row in data_rows[:max_rows]:
                record = {}
                for j, val in enumerate(row):
                    key = headers[j] if j < len(headers) else f"col_{j}"
                    record[key] = val
                records.append(record)

            logger.info("read_csv_success", file_path=file_path, rows=total, encoding=detected_encoding)
            return ToolResult(
                success=True,
                data={
                    "headers": headers,
                    "rows": records,
                    "total_rows": total,
                    "truncated": total > max_rows,
                    "detected_encoding": detected_encoding,
                    "detected_delimiter": repr(detected_delimiter),
                },
            )

        except ValueError as exc:
            return ToolResult(success=False, error=str(exc))
        except Exception as exc:
            logger.error("read_csv_failed", file_path=file_path, error=str(exc))
            return ToolResult(success=False, error=f"Failed to read CSV: {exc}")


# ── read_code ─────────────────────────────────────────────────────────────────


class ReadCodeTool(BaseTool):
    """Read source code with line numbers and approximate structure extraction."""

    name = "read_code"
    description = (
        "Read a source code file. Returns content with line numbers, detected "
        "language, line count, and approximate structure extraction (functions, "
        "classes, imports via regex — not a full parser). Supports any text-based "
        "source file."
    )
    parameters = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the source file.",
            },
            "line_range": {
                "type": "string",
                "description": "Line range to read, e.g. '1-50' or '100-200'. Default: entire file.",
                "default": "",
            },
            "extract_structure": {
                "type": "boolean",
                "description": "Extract functions, classes, imports (regex-based, approximate). Default: true.",
                "default": True,
            },
            "max_lines": {
                "type": "integer",
                "description": "Max lines to return. Default: 2000.",
                "default": 2000,
            },
        },
        "required": ["file_path"],
    }

    _LANG_MAP = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".tsx": "tsx",
        ".jsx": "jsx",
        ".java": "java",
        ".go": "go",
        ".rs": "rust",
        ".cpp": "cpp",
        ".c": "c",
        ".h": "c",
        ".cs": "csharp",
        ".rb": "ruby",
        ".php": "php",
        ".swift": "swift",
        ".kt": "kotlin",
        ".scala": "scala",
        ".sh": "bash",
        ".sql": "sql",
        ".html": "html",
        ".css": "css",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".toml": "toml",
        ".json": "json",
        ".xml": "xml",
        ".r": "r",
        ".R": "r",
        ".md": "markdown",
        ".vue": "vue",
        ".svelte": "svelte",
    }

    # Regex patterns for *approximate* structure extraction.
    # These are heuristics, not parsers — they'll miss edge cases.
    _STRUCTURE_PATTERNS: dict[str, dict[str, str]] = {
        "python": {
            "functions": r"^(?:async\s+)?def\s+(\w+)\s*\(",
            "classes": r"^class\s+(\w+)",
            "imports": r"^(?:from\s+\S+\s+)?import\s+.+",
        },
        "javascript": {
            "functions": r"(?:export\s+)?(?:async\s+)?function\s+(\w+)|(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s+)?\(?.*?\)?\s*=>",
            "classes": r"(?:export\s+)?class\s+(\w+)",
            "imports": r"^import\s+.+",
        },
        "typescript": {
            "functions": r"(?:export\s+)?(?:async\s+)?function\s+(\w+)|(?:const|let)\s+(\w+)\s*(?::\s*\S+\s*)?=\s*(?:async\s+)?\(?.*?\)?\s*=>",
            "classes": r"(?:export\s+)?(?:abstract\s+)?class\s+(\w+)",
            "imports": r"^import\s+.+",
        },
        "go": {
            "functions": r"^func\s+(?:\(\w+\s+\*?\w+\)\s+)?(\w+)\s*\(",
            "classes": r"^type\s+(\w+)\s+struct",
            "imports": r'^\s*"[^"]+"|^\s*\w+\s+"[^"]+"',
        },
        "rust": {
            "functions": r"(?:pub\s+)?(?:async\s+)?fn\s+(\w+)",
            "classes": r"(?:pub\s+)?(?:struct|enum|trait)\s+(\w+)",
            "imports": r"^use\s+.+",
        },
        "java": {
            "functions": r"(?:public|private|protected)?\s*(?:static\s+)?(?:\w+\s+)+(\w+)\s*\([^)]*\)\s*\{",
            "classes": r"(?:public\s+)?(?:abstract\s+)?class\s+(\w+)",
            "imports": r"^import\s+.+",
        },
    }

    def __init__(self, output_dir: str = "./hydra_output", state_manager: "StateManager | None" = None) -> None:
        self._output_dir = output_dir
        self._state_manager = state_manager

    async def execute(
        self,
        file_path: str,
        line_range: str = "",
        extract_structure: bool = True,
        max_lines: int = 2000,
    ) -> ToolResult:
        try:
            path = safe_read_path(file_path, allowed_roots=[self._output_dir, Path.cwd()])
            suffix = path.suffix.lower()
            language = self._LANG_MAP.get(suffix, "unknown")

            # Read with encoding fallback
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                content = path.read_text(encoding="latin-1")

            all_lines = content.splitlines()
            total_lines = len(all_lines)

            # Apply line range
            offset = 1
            display_lines = all_lines
            if line_range:
                parts = line_range.split("-")
                start = max(1, int(parts[0]))
                end = min(total_lines, int(parts[1])) if len(parts) > 1 else total_lines
                display_lines = all_lines[start - 1 : end]
                offset = start

            # Truncate
            truncated = len(display_lines) > max_lines
            display_lines = display_lines[:max_lines]

            # Format with line numbers
            numbered = [f"{i + offset:>4} | {line}" for i, line in enumerate(display_lines)]

            result: dict = {
                "language": language,
                "total_lines": total_lines,
                "content": "\n".join(numbered),
                "truncated": truncated,
                "file_size_bytes": path.stat().st_size,
            }

            # Approximate structure extraction
            if extract_structure and language in self._STRUCTURE_PATTERNS:
                patterns = self._STRUCTURE_PATTERNS[language]
                structure: dict[str, list] = {}
                for kind, pattern in patterns.items():
                    matches = []
                    for i, line in enumerate(all_lines, 1):
                        m = re.search(pattern, line)
                        if m:
                            name = next((g for g in m.groups() if g), line.strip())
                            matches.append({"name": name, "line": i})
                    structure[kind] = matches
                result["structure"] = structure

            logger.info("read_code_success", file_path=file_path, language=language, lines=total_lines)
            return ToolResult(success=True, data=result)

        except ValueError as exc:
            return ToolResult(success=False, error=str(exc))
        except Exception as exc:
            logger.error("read_code_failed", file_path=file_path, error=str(exc))
            return ToolResult(success=False, error=f"Failed to read code file: {exc}")
